from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "professor-brief.docx"

INK = "0B2545"
BLUE = "2E74B5"
PALE = "E8EEF5"
LIGHT = "F4F6F9"
MUTED = "566270"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def mark_first_row_as_header(table):
    """Expose the first row as a header to Word and assistive technology."""
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_cell_border(cell, color="D9DEE5", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def style_run(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_rich_paragraph(doc, pieces, *, style=None, before=0, after=6, line=1.1, keep=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep
    for text, bold, italic, color in pieces:
        style_run(p.add_run(text), bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.82)
section.right_margin = Inches(0.82)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.08

for name, size, color, before, after in (
    ("Heading 1", 15, BLUE, 12, 5),
    ("Heading 2", 12, BLUE, 8, 3),
    ("Heading 3", 11, INK, 6, 2),
):
    style = styles[name]
    style.font.name = "Aptos Display" if name == "Heading 1" else "Aptos"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
style_run(header.add_run("RESEARCH PROSPECTUS  |  REDE"), size=8, bold=True, color=MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_run(footer.add_run("Jaya Krishna J  |  24 August 2026  |  Executable baseline; theorem and multi-rater validation proposed"), size=7.5, color=MUTED)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(2)
style_run(title.add_run("Decision-safe auto-evaluation"), size=23, bold=True, color=INK)
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(8)
style_run(subtitle.add_run("of robot policies"), size=23, bold=True, color=BLUE)
meta = doc.add_paragraph()
meta.paragraph_format.space_after = Pt(10)
style_run(meta.add_run("Jaya Krishna J  ·  thesis research prospectus  ·  Validate stage"), size=9.5, bold=True, color=MUTED)

callout = doc.add_table(rows=1, cols=1)
callout.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(callout, [9360])
cell = callout.cell(0, 0)
set_cell_shading(cell, PALE)
set_cell_border(cell, color=PALE)
p = cell.paragraphs[0]
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.08
style_run(p.add_run("Research question  "), size=10.5, bold=True, color=BLUE)
style_run(p.add_run("Can a fail-safe policy-decision certificate reduce human review with VLM judgments while preserving simultaneous error control—and abstain to human-only inference when the evaluator cannot help?"), size=10.5, color=INK)

add_heading(doc, "Why this is a research problem", 1)
add_rich_paragraph(doc, [("Episode accuracy does not establish policy-decision fidelity. Structured VLM errors can cancel for one comparison and amplify for another. The thesis therefore treats automated judgment as optional evidence, produces simultaneous superiority/equivalence decisions, and returns to human-only inference when a frozen gate fails.", False, False, INK)])

add_heading(doc, "Preliminary evidence", 1)
evidence = doc.add_table(rows=1, cols=3)
evidence.alignment = WD_TABLE_ALIGNMENT.LEFT
headers = ["Matched human metrics", "Public VLM predictions", "Conditional label value"]
for idx, text in enumerate(headers):
    c = evidence.rows[0].cells[idx]
    set_cell_shading(c, INK)
    set_cell_border(c, color=INK)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(c.paragraphs[0].add_run(text), size=9, bold=True, color=WHITE)
row = evidence.add_row()
values = [
    ("3,879 sessions\n28 policy pairs\n13.0-point mean shift", "Binary success was saturated: both policies failed in 81.8% of sessions."),
    ("25 released runs\n158 matched sessions\n24 complete models", "Every complete model changed at least one descriptive policy decision."),
    ("5 of 25 runs helped\n~14-15% best-case gain\n~25% harmful example", "Fixed PPI helps only when the evaluator reduces residual variance."),
]
for idx, (lead, detail) in enumerate(values):
    c = row.cells[idx]
    set_cell_shading(c, LIGHT)
    set_cell_border(c)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line_idx, line_text in enumerate(lead.split("\n")):
        if line_idx:
            p.add_run("\n")
        style_run(p.add_run(line_text), size=11, bold=True, color=BLUE)
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_after = Pt(0)
    style_run(p2.add_run(detail), size=8.5, color=INK)
set_table_geometry(evidence, [3120, 3120, 3120])

note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(4)
note.paragraph_format.space_after = Pt(4)
style_run(note.add_run("Evidence status: "), size=8.5, bold=True, color=MUTED)
style_run(note.add_run("exploratory. Released test artifacts informed method selection; the strengthened design separates evaluator audit from inference and adds new multi-rater validation."), size=8.5, italic=True, color=MUTED)

add_heading(doc, "Defensible contribution", 1)
add_rich_paragraph(doc, [("A verified conservative certificate plus a research-hard graph-aware validity/no-harm target.", True, False, INK)])
add_rich_paragraph(doc, [("The split-sample Bonferroni baseline now runs and fails closed. For an MASc/MSc, the contribution is its locked multi-rater and shift validation. For a PhD, the thesis must prove cross-fitted simultaneous validity and a high-probability no-harm bound; PPI plus a gate is not enough.", False, False, INK)], after=2)

doc.add_page_break()

add_heading(doc, "Executable baseline and thesis risk", 1)
add_rich_paragraph(doc, [("The new roboeval certify command freezes the proxy frame, rejects audit/inference leakage and missing graph edges, gates residual variance, bias, and supported shift strata, then applies fixed PPI or the identical human-only fallback. Its Bonferroni-normal intervals are a transparent baseline—not the proposed graph-aware theorem.", False, False, INK)])

gates = doc.add_table(rows=1, cols=3)
gates.alignment = WD_TABLE_ALIGNMENT.LEFT
for idx, text in enumerate(("Coverage gate", "Cost gate", "Abstention rule")):
    c = gates.rows[0].cells[idx]
    set_cell_shading(c, PALE)
    set_cell_border(c)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(c.paragraphs[0].add_run(text), size=9.5, bold=True, color=BLUE)
row = gates.add_row()
for idx, text in enumerate(("Simultaneous 95% target", "Review-time ratio below frozen threshold", "Human-only if any gate fails")):
    c = row.cells[idx]
    set_cell_border(c)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(c.paragraphs[0].add_run(text), size=9.5, bold=True if idx == 2 else False, color=INK)
set_table_geometry(gates, [3120, 3120, 3120])

add_heading(doc, "Two-track validation", 1)
add_rich_paragraph(doc, [("The minimum track uses stable public episodes plus new blinded ratings from at least three raters per selected episode; it supports internal validation without waiting for a future release. A prospective lab frame with new tasks, policies, sites, or robot conditions is the preferred external-validation track.", False, False, INK)], after=2)

add_heading(doc, "Research assets already built", 1)
assets = [
    "Deterministic matched-policy inference and finite-population PPI simulation.",
    "Runnable graph certificate with helpful, harmful, shift, leakage, and missing-edge tests.",
    "Reference-aware calibration, two-phase adjudication, and fail-closed evidence checks.",
    "A documented CLI, public-artifact fetchers, hashed manifests, and data-boundary enforcement.",
]
for item in assets:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    style_run(p.add_run(item), size=9.5, color=INK)

add_heading(doc, "Twelve-month plan", 1)
roadmap = doc.add_table(rows=1, cols=3)
roadmap.alignment = WD_TABLE_ALIGNMENT.LEFT
for idx, text in enumerate(("Months", "Research output", "Decision gate")):
    c = roadmap.rows[0].cells[idx]
    set_cell_shading(c, INK)
    set_cell_border(c, color=INK)
    style_run(c.paragraphs[0].add_run(text), size=8.5, bold=True, color=WHITE)
for values in (
    ("1-2", "Graph, margin, theorem target, simulation, ethics/data decision", "Coverage survives adversarial proxies"),
    ("3-4", "Blinded multi-rater pilot and audit/inference split", "Reliability, cost, and support pass"),
    ("5-7", "Locked comparison + graph-aware proof attempt", "Validity/no-harm claim survives"),
    ("8-9", "Task/site shift or prospective lab validation", "Internal vs external claims separated"),
    ("10-12", "Thesis, protocol archive, reproducibility release", "Claim/privacy/accessibility pass"),
):
    row = roadmap.add_row()
    for idx, text in enumerate(values):
        c = row.cells[idx]
        set_cell_border(c)
        if int(values[0].split("-")[0]) % 2 == 1:
            set_cell_shading(c, LIGHT)
        style_run(c.paragraphs[0].add_run(text), size=8.3, bold=idx == 0, color=INK)
set_table_geometry(roadmap, [900, 4770, 3690])

add_heading(doc, "The supervision decision", 1)
ask = doc.add_table(rows=1, cols=1)
ask.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(ask, [9360])
c = ask.cell(0, 0)
set_cell_shading(c, PALE)
set_cell_border(c, color=PALE)
p = c.paragraphs[0]
p.paragraph_format.space_after = Pt(0)
style_run(p.add_run("Is cross-fitted graph-aware validity/no-harm a thesis-worthy core for this lab—and which prospective evaluation frame would make the robotics claim decisive?"), size=10.5, bold=True, color=INK)

refs = doc.add_paragraph()
refs.paragraph_format.space_before = Pt(5)
refs.paragraph_format.space_after = Pt(0)
refs.paragraph_format.line_spacing = 1.0
style_run(refs.add_run("Selected sources: "), size=7.5, bold=True, color=MUTED)
style_run(refs.add_run("RoboArena; RoboReward; PPI/PPI++; SureSim; Post-hoc Regression; MultiPPI; AM-PPI; Prediction-Powered E-Values. Full novelty and claim boundaries: docs/thesis-proposal.md."), size=7.5, color=MUTED)

doc.core_properties.title = "Decision-safe auto-evaluation of robot policies"
doc.core_properties.subject = "Professor-facing thesis research prospectus"
doc.core_properties.author = "Jaya Krishna J"
doc.core_properties.keywords = "robot policy evaluation, VLM, prediction-powered inference, thesis proposal"
for table in doc.tables:
    mark_first_row_as_header(table)
doc.save(OUTPUT)
print(OUTPUT)
